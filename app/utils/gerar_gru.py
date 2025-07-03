from playwright.sync_api import sync_playwright, expect
import time
from datetime import date, datetime, timedelta
import requests
import os
from docx import Document
from docx.shared import Inches
import win32com.client
import pythoncom
import unicodedata


def gerar_pdf(usuario, cnpj, nome, ref, valor_total, qr_txt):

    documento = Document(
        f"C:\\Users\\{usuario}\\Documents\\GRU_PIX_MODELO.docx")
    qr_code = f"C:\\Users\\{usuario}\\Documents\\qr_code.png"

    data_atual = date.today()
    data_hora_atual = datetime.now().strftime("%d/%m/%Y %H:%M")

    referencias = {
        'NOME_CHAVE': nome,
        'CNPJ_CHAVE': cnpj,
        'REF_CHAVE': ref,
        'VALOR_CHAVE': valor_total,
        'DATA_ATUAL_CHAVE': f'{data_atual.day:02}/{data_atual.month:02}/{data_atual.year}',
        'CODIGO_QR_TXT': qr_txt,
        'DATA_HORA_CHAVE': data_hora_atual
    }

    for paragrafo in documento.paragraphs:
        for run in paragrafo.runs:
            for ref, valor in referencias.items():
                if ref in run.text:
                    # Substitui o código de referência pelo valor correspondente
                    run.text = run.text.replace(ref, valor)
                if "IMG_QR" in run.text:
                    texto_restante = paragrafo.text.replace("IMG_QR", "")
                    paragrafo.text = texto_restante

                    # Inserir a imagem após o texto atualizado
                    run = paragrafo.add_run()
                    run.add_picture(qr_code, width=Inches(2))

    os.remove(qr_code)

    nome_arquivo = nome.split(" ")[0]
    nome_arquivo = f"{nome_arquivo}_{valor_total}_grupix"
    nome_arquivo = nome_arquivo.replace(" ", "")

    documento.save(f"C:\\Users\\{usuario}\\Documents\\{nome_arquivo}.docx")

    pythoncom.CoInitialize()

    time.sleep(2)

    word = win32com.client.Dispatch("Word.Application")

    doc = word.Documents.Open(
        f"C:\\Users\\{usuario}\\Documents\\{nome_arquivo}.docx")
    # 17 é o formato PDF
    doc.SaveAs(
        f"C:\\Users\\{usuario}\\Documents\\{nome_arquivo}.pdf", FileFormat=17)
    doc.Close()
    word.Quit()

    os.remove(f"C:\\Users\\{usuario}\\Documents\\{nome_arquivo}.docx")


def gru_pix(page, usuario, cnpj, nome, valor):

    try:
        page.goto("https://pagtesouro.tesouro.gov.br/portal-gru/#/pagamento-gru/formulario?servico=000034",
                  wait_until='domcontentloaded')

        expect(page.locator(
            "input[placeholder='Digite o CPF ou CNPJ do contribuinte']")).to_be_visible(timeout=30000)
        page.locator(
            "input[placeholder='Digite o CPF ou CNPJ do contribuinte']").click()
        page.locator("input[placeholder='Digite o CPF ou CNPJ do contribuinte']").type(
            cnpj, delay=50)

        page.locator(
            "input[placeholder='Digite o nome do contribuinte ou a razão social da empresa']").fill(nome)

        if len(cnpj) > 14:
            num_ref = "265"
        else:
            num_ref = "215"

        page.locator(
            "input[placeholder='Digite a referência do serviço, caso o serviço solicite']").fill(num_ref)

        competencia = f"{date.today().month}/{date.today().year}"
        page.locator("input[placeholder='Digite o mês e o ano da competência']").fill(
            competencia)

        vencimento = date.today() + timedelta(days=20)
        vencimento = f"{vencimento.day}/{vencimento.month}/{vencimento.year}"
        page.locator("input[placeholder='Digite a data de vencimento']").fill(
            vencimento)

        page.locator(
            "input[placeholder='Digite o valor principal do pagamento']").click()
        page.locator(
            "input[placeholder='Digite o valor principal do pagamento']").fill(str(valor))

        estado = "Ocorreu algum erro ao preencher o formulário."

        page.locator("button[type='submit']").click()

        page.frame_locator("#iFrameResizer0").locator(
            ".meio-pagamento >> nth = 0").click()
        page.frame_locator("#iFrameResizer0").locator("#btnPgto").click()
        time.sleep(2)

        qr_txt = page.frame_locator("#iFrameResizer0").locator(
            ".noselect").inner_text()
        element = page.frame_locator("#iFrameResizer0").locator(
            ".qr-code-img >> nth = 1")  # Localizador do elemento
        element.screenshot(
            path=f"C:\\Users\\{usuario}\\Documents\\qr_code.png")

        gerar_pdf(usuario, cnpj, nome, num_ref, valor, qr_txt)

        return "Sucesso"

    except Exception as e:
        print(e)
        return estado


def remover_acentos(texto):
    return ''.join(c for c in unicodedata.normalize('NFKD', texto) if unicodedata.category(c) != 'Mn')

def gru_simples(usuario, cnpj, nome, valor):
    if len(cnpj) > 14:
        num_ref = "265"
    else:
        num_ref = "215"
    nome = remover_acentos(nome)
    nome_arquivo = nome.split(" ")[0]
    nome_arquivo = f"{nome_arquivo}_R${valor}_grusimples"
    nome_arquivo = nome_arquivo.replace(" ", "")

    nome = nome.replace(" ","+").replace(".","")
    competencia = f"{date.today().month:02}%2F{date.today().year}"
    vencimento = date.today() + timedelta(days=20)
    vencimento = f"{vencimento.day:02}%2F{vencimento.month:02}%2F{vencimento.year}"

    valor= str(valor).replace(".","%2C").replace(",","%2C")
    url = f"https://pagtesouro.tesouro.gov.br/api/gru/portal/boleto-gru?codigoUg=010001&codigoRecolhimento=98815-4&cpfCnpjContribuinte={cnpj}&nomeContribuinte={nome}&numeroReferencia={num_ref}&competencia={competencia}&vencimento={vencimento}&valorPrincipal={valor}"

    caminho_pdf = f"C:\\Users\\{usuario}\\Documents\\{nome_arquivo}.pdf"
    
    response = requests.get(url)
    
    if response.status_code == 200:
        with open(caminho_pdf, 'wb') as file:
            file.write(response.content)
        return "Sucesso"
    else:
        return(f"Erro ao baixar o PDF. Código de status HTTP: {response.status_code}")
    

def run_gerar_gru(cnpj, nome, valor, tipo):
    usuario = os.getlogin().lower()

    if tipo == "simples":
        estado = gru_simples(usuario, cnpj, nome, valor)
    elif tipo == "pix":

        with sync_playwright() as p:

            browser = p.chromium.launch(
                channel="chromium", headless=True, downloads_path=f"C:\\Users\\{usuario}\\Documents")
            page = browser.new_page()
            page.set_default_timeout(90000)

            estado = gru_pix(page, usuario, cnpj, nome, valor)

            page.close()
            browser.close()

    return estado


    